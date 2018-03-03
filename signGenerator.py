#we use this library to deal with excel; it's quite capable and grabs the data quite well.
import openpyxl
#this library is for creating word documents; we will need this
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
#we need the sys library to grab command line arguments for the file; then we need the os library to open various folders
import os, sys
#import any functions from misc.py here as if they were in the sheet. 
from misc import *

#the intent of the program is to accept command line arguments for a filename for the excel sheet to load so it would run like signGenerator.py "Human Library Books.xlsx" 
#therefor; we need to check if it has an argument 
if(len(sys.argv)>1):
    # getcwd() gets the current path; os.sep is a separator between directories and sys.argv[1]  grabs from the argument array list given by whatever you enter after the program name
    filePath = os.getcwd() + os.sep + sys.argv[1]
    #load the excel sheet into a workbook object
    workbook = openpyxl.load_workbook(filePath)
    #grab the pages in an array of the excel book in an array
    workbookPages = workbook.get_sheet_names()
    #instantiate lists for both the question and title pages.
    titlePageList = []
    questionPageList = []
    #loop through them
    for i in range(0, len(workbookPages)):
        #grab the sheet from the list of pages
        workbookSheet = workbook[workbookPages[i]]
        #if it's the title page; we go through it
        if workbookPages[i] == "Title page":
            #we start at 4 because the data starts on the 4th row; we go till no more rows
            for y in range(4, workbookSheet.max_row+1):
                #store needed values in the sheet for printing
                tmpList = []
                #here we go in range from the start to the last column of the excel sheet
                for x in range(0, workbookSheet.max_column):
                    #if there's a value returnable; we add it to a temp list
                    if workbookSheet[rowColConvert(y,x)].value:
                        tmpList.append(workbookSheet[rowColConvert(y,x)].value)
                    #if there's not; we don't and instead we fill it with a blank value
                    else:
                        tmpList.append(" ")
                #once the tmp list is filled for the row we add it to the title page list so we can further loop through it
                titlePageList.append(tmpList)
        #if it's the questions page; we go through it
        if workbookPages[i] == "Question page":
            for y in range(3, workbookSheet.max_row+1):
                tmpList = []
                for x in range(0, workbookSheet.max_column):
                    if workbookSheet[rowColConvert(y,x)].value:
                        tmpList.append(workbookSheet[rowColConvert(y,x)].value)
                    else:
                        tmpList.append(" ")
                questionPageList.append(tmpList)
    for x in range(0, len(titlePageList)):
        #initiate word document for editing
        document = Document()
        #here we create styles for text formatting. 
        style = document.styles['Title']
        font = style.font
        font.name = 'Oswald SemiBold'
        font.size = Pt(65)
        styles = document.styles
        style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Source Serif Pro'
        font.size = Pt(18)
        styles = document.styles
        style = styles.add_style('QuestionHeader', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.bold = True
        font.name = 'Open Sans'
        font.size = Pt(18)
        styles = document.styles
        style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(18)
        #add stuff from the list to the pages line by line using the previous styles
        title = document.add_paragraph(titlePageList[x][1])
        title.style=document.styles['Title']
        author = document.add_paragraph("By: " + str(titlePageList[x][0]))
        author.style=document.styles['Name']
        new_section=document.add_section(WD_SECTION.ODD_PAGE)
        qh1 = document.add_paragraph("Opening Question")
        qh1.style=document.styles['QuestionHeader']
        q1 = document.add_paragraph(titlePageList[x][2])
        q1.style=document.styles['Question']
        qh2 = document.add_paragraph("Need Help?")
        qh2.style=document.styles['QuestionHeader']
        q2 = document.add_paragraph(titlePageList[x][3])
        q2.style=document.styles['Question']
        qh3 = document.add_paragraph("End Discussion Suggestion")
        qh3.style=document.styles['QuestionHeader']
        q3 = document.add_paragraph("I appreciate your interest in my topic. However, I think it is best that we end our discussion now and agree to disagree")
        q3.style=document.styles['Question']
        qh4 = document.add_paragraph("Reminder")
        qh4.style=document.styles['QuestionHeader']
        q4 = document.add_paragraph(titlePageList[x][5])
        q4.style=document.styles['Question']
        document.save(str(os.getcwd()) + os.sep + "wordFiles" + os.sep + str(titlePageList[x][0])+ "-Title.docx")


    for x in range(0, len(questionPageList)):
        #initiate Document 
        document = Document()
        #Here we create styles for text formatting
        style = document.styles['Title']
        font = style.font
        font.name = 'Oswald SemiBold'
        font.size = Pt(65)
        styles = document.styles
        style = styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Source Serif Pro'
        font.size = Pt(18)
        styles = document.styles
        style = styles.add_style('QuestionHeader', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.bold = True
        font.name = 'Open Sans'
        font.size = Pt(18)
        styles = document.styles
        style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Open Sans'
        font.size = Pt(16)

        #add all questions
        name = document.add_paragraph(questionPageList[x][0])
        name.style = document.styles['Question']
        title = document.add_paragraph(questionPageList[x][1])
        title.style=document.styles['Question']
        qPart = document.add_paragraph("Question suggestions for your book:")
        qPart.style=document.styles['QuestionHeader']
        for y in range(0, len(questionPageList[x])):
            q1 = document.add_paragraph(questionPageList[x][y])
            q1.style=document.styles['Question']
        document.save(str(os.getcwd()) + os.sep + "wordFiles" + os.sep + str(questionPageList[x][0]) + "-Questions.docx")


        


